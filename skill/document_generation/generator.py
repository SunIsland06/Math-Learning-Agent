"""
文档生成技能 —— 将 Markdown 内容转换为 PDF、Word 或 Markdown 文档。

特性：
- PDF 生成：使用 fpdf2，支持中文、页眉页脚、目录、表格
- Word 生成：使用 python-docx，支持标题层级、段落格式
- Markdown 生成：直接保存 Markdown 文本
- 数学公式保留 LaTeX 表示（pdf/word 中显示为 $$...$$ 代码）
- 自动保存到用户会话目录，返回下载 URL

入口函数：skill_entry(payload) → str
"""

import re
import sys
import os
from pathlib import Path
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple


def skill_entry(payload: dict) -> str:
    """技能入口 —— 根据 Markdown 内容生成文档文件。"""
    payload = payload or {}
    params = payload.get("params") or {}
    content = params.get("content") or payload.get("input", "")
    fmt = (params.get("format") or "pdf").lower()
    title = params.get("title", "学习报告")
    author = params.get("author", "")
    header_text = params.get("header_text", "")
    footer_text = params.get("footer_text", "")
    include_toc = params.get("include_toc", True)
    output_dir = params.get("output_dir")
    output_url_prefix = params.get("output_url_prefix", "")

    if not content or not content.strip():
        return "错误：请提供要生成的文档内容（Markdown 格式）。"

    content = content.strip()

    if output_dir:
        folder = Path(str(output_dir))
    else:
        folder = Path(__file__).resolve().parents[2] / "userdata" / "_generated"
    folder.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S%f")
    safe_title = re.sub(r"[^\w一-鿿_-]", "_", title)[:30]

    try:
        if fmt == "docx":
            filepath = _generate_docx(content, title, author, folder, safe_title, timestamp)
        elif fmt == "md":
            filepath = _generate_markdown(content, folder, safe_title, timestamp)
        else:
            filepath = _generate_pdf(content, title, author, header_text, footer_text,
                                     include_toc, folder, safe_title, timestamp)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return f"【文档生成失败】{fmt} 格式生成时出错: {e}"

    filename = filepath.name
    url = f"{output_url_prefix}{filename}" if output_url_prefix else f"/userdata/_generated/{filename}"

    return (
        f"【文档已生成】\n\n"
        f"标题: {title}\n"
        f"格式: {fmt.upper()}\n"
        f"文件名: {filename}\n\n"
        f"[点击下载文档]({url})\n\n"
        f"下载链接: `{url}`"
    )


# ============================================================
# 跨平台字体查找
# ============================================================

def _find_cjk_font_paths() -> List[str]:
    """跨平台搜索 CJK 中文字体文件路径。

    按优先级搜索 Windows、Linux、macOS 常见字体目录，
    返回找到的 TrueType/OpenType 字体文件路径列表。
    """
    candidates = []

    if sys.platform == "win32":
        windir = os.environ.get("WINDIR", r"C:\Windows")
        fonts_dir = os.path.join(windir, "Fonts")
        candidates = [
            os.path.join(fonts_dir, "msyh.ttc"),      # 微软雅黑
            os.path.join(fonts_dir, "msyhbd.ttc"),    # 微软雅黑 粗体
            os.path.join(fonts_dir, "msyh.ttf"),      # 微软雅黑 (ttf)
            os.path.join(fonts_dir, "simsun.ttc"),    # 宋体
            os.path.join(fonts_dir, "simhei.ttf"),    # 黑体
            os.path.join(fonts_dir, "simkai.ttf"),    # 楷体
            os.path.join(fonts_dir, "mingliu.ttc"),   # 细明体
            os.path.join(fonts_dir, "msgothic.ttc"),  # MS Gothic
        ]
    elif sys.platform == "darwin":
        candidates = [
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/System/Library/Fonts/STHeiti Medium.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
            "/System/Library/Fonts/AppleSDGothicNeo.ttc",
        ]
    else:  # Linux
        candidates = [
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
            "/usr/share/fonts/truetype/arphic/uming.ttc",
            "/usr/share/fonts/truetype/arphic/ukai.ttc",
        ]

    return [p for p in candidates if os.path.isfile(p)]


def _register_cjk_font(pdf) -> str:
    """注册 CJK 中文字体到 fpdf2，失败时创建内置小字体的子集。

    返回注册成功的字体名称，若全部失败则回退到 Courier。
    """
    font_paths = _find_cjk_font_paths()

    for fp in font_paths:
        try:
            pdf.add_font("CJK", "", fp, uni=True)
            try:
                pdf.add_font("CJK", "B", fp, uni=True)
            except Exception:
                pass
            try:
                # CJK 字体通常没有真正斜体，用常规字体替代 I 样式
                pdf.add_font("CJK", "I", fp, uni=True)
            except Exception:
                pass
            return "CJK"
        except Exception:
            continue

    # 最终回退：使用 Courier（仅支持 ASCII，中文会丢失）
    return "Courier"


# ============================================================
# PDF 生成
# ============================================================

def _generate_pdf(content: str, title: str, author: str,
                  header_text: str, footer_text: str, include_toc: bool,
                  folder: Path, safe_title: str, timestamp: str) -> Path:
    """使用 fpdf2 生成 PDF 文档。"""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_left_margin(15)
    pdf.set_right_margin(15)

    font_name = _register_cjk_font(pdf)
    effective_width = pdf.w - pdf.l_margin - pdf.r_margin

    # ---- 封面 ----
    pdf.add_page()
    try:
        if font_name != "Courier":
            pdf.set_font(font_name, "B", 22)
        else:
            pdf.set_font(font_name, "B", 18)
        pdf.ln(40)
        pdf.multi_cell(effective_width, 14, title, align="C")
        pdf.ln(10)

        if author:
            pdf.set_font(font_name, "", 12)
            pdf.cell(effective_width, 10, f"作者: {author}", align="C")
            pdf.ln(10)

        pdf.set_font(font_name, "", 10)
        pdf.cell(effective_width, 10, f"生成时间: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC", align="C")
        pdf.ln(10)
        pdf.cell(effective_width, 10, "由 Math-Learning-Agent (MLA) 数智教师生成", align="C")

        if font_name == "Courier":
            pdf.ln(16)
            pdf.set_font(font_name, "", 9)
            pdf.multi_cell(effective_width, 6,
                           "(提示: 服务器未安装中文字体，中文内容可能无法正确显示。"
                           "建议在系统中安装中文字体后重试。)")
    except Exception:
        pdf.set_font("Courier", "", 11)
        pdf.cell(effective_width, 10, f"Title: {title}", align="C")

    # ---- 目录 ----
    if include_toc:
        toc_items = _extract_toc(content)
        if toc_items:
            pdf.add_page()
            try:
                pdf.set_font(font_name, "B", 16)
                pdf.cell(effective_width, 12, "目录", align="C")
                pdf.ln(14)
                pdf.set_font(font_name, "", 11)
                for level, heading, _ in toc_items:
                    indent = "    " * (level - 1)
                    prefix = {1: "● ", 2: "  - ", 3: "    · "}.get(level, "    · ")
                    line = f"{indent}{prefix}{heading}"
                    try:
                        pdf.cell(effective_width, 8, line)
                        pdf.ln(8)
                    except Exception:
                        pdf.ln(4)
            except Exception:
                pass

    # ---- 正文 ----
    pdf.add_page()

    if header_text:
        try:
            pdf.header = lambda p=pdf, f=font_name, h=header_text: _pdf_header(p, f, h)
        except Exception:
            pass
    if footer_text:
        try:
            pdf.footer = lambda p=pdf, f=font_name, ft=footer_text: _pdf_footer(p, f, ft)
        except Exception:
            pass

    _render_markdown_to_pdf(pdf, content, font_name)

    filepath = folder / f"{safe_title}_{timestamp}.pdf"
    pdf.output(str(filepath))
    return filepath


def _pdf_header(pdf, font_name, header_text):
    """PDF 页眉（居中文字）。"""
    pdf.set_font(font_name, "", 8)
    pdf.cell(0, 8, header_text, align="C")
    pdf.ln(10)


def _pdf_footer(pdf, font_name, footer_text):
    """PDF 页脚（居中文字 + 页码）。"""
    pdf.set_y(-15)
    pdf.set_font(font_name, "", 8)
    pdf.cell(0, 10, f"{footer_text}  |  第 {pdf.page_no()} 页", align="C")


def _extract_toc(content: str) -> List[Tuple[int, str, int]]:
    """从 Markdown 内容提取目录项。"""
    toc = []
    for line in content.splitlines():
        match = re.match(r"^(#{1,3})\s+(.+)", line)
        if match:
            level = len(match.group(1))
            heading = match.group(2).strip()
            heading = re.sub(r"\$\$?(.+?)\$\$?", r"\1", heading)
            heading = re.sub(r"[\\[\]{}]", "", heading)
            toc.append((level, heading[:60], 0))
    return toc


def _render_markdown_to_pdf(pdf, content: str, font_name: str):
    """将 Markdown 内容逐行渲染为 PDF。

    支持的 Markdown 元素：
    - 标题（h1-h4）
    - 代码块（带灰底）
    - 无序列表
    - 水平分割线
    - 表格（| 分隔的简单表格）
    - LaTeX 公式（保留 $...$ 和 $$...$$ 显示）
    - 粗体/斜体（保留 * 标记）
    - 链接（保留 [text](url) 格式）
    """
    effective_width = pdf.w - pdf.l_margin - pdf.r_margin - 4
    lines = content.splitlines()
    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        # ---- 代码块 ----
        if line.strip().startswith("```"):
            if in_code_block:
                _render_code_block(pdf, code_buffer, effective_width)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # ---- 表格 ----
        if "|" in line and line.strip().startswith("|"):
            table_lines = [line]
            j = i + 1
            while j < len(lines) and "|" in lines[j] and lines[j].strip().startswith("|"):
                table_lines.append(lines[j])
                j += 1
            # 跳过分隔行（如 |---|---|）
            filtered = [tl for tl in table_lines if not re.match(r"^\|[\s\-:|]+\|$", tl.strip())]
            if len(filtered) >= 2:
                _render_table(pdf, font_name, filtered, effective_width)
            else:
                # 非表格，按普通段落处理
                _render_paragraph(pdf, font_name, line, effective_width)
            i = j
            continue

        # ---- 标题 ----
        h_match = re.match(r"^(#{1,4})\s+(.+)", line)
        if h_match:
            level = len(h_match.group(1))
            heading = h_match.group(2).strip()
            _pdf_section_heading(pdf, font_name, heading, level)
            i += 1
            continue

        # ---- 无序列表 ----
        if re.match(r"^\s*[-*+]\s+", line):
            _render_list_item(pdf, font_name, line, effective_width)
            i += 1
            continue

        # ---- 有序列表 ----
        if re.match(r"^\s*\d+[.)]\s+", line):
            _render_list_item(pdf, font_name, line, effective_width)
            i += 1
            continue

        # ---- 水平线 ----
        if re.match(r"^[-*_]{3,}\s*$", line):
            pdf.ln(4)
            pdf.cell(effective_width, 0, "", border="T")
            pdf.ln(4)
            i += 1
            continue

        # ---- 空行 ----
        if not line.strip():
            pdf.ln(4)
            i += 1
            continue

        # ---- 引用块 ----
        if line.strip().startswith("> "):
            _render_blockquote(pdf, font_name, line, effective_width)
            i += 1
            continue

        # ---- 普通段落 ----
        _render_paragraph(pdf, font_name, line, effective_width)
        i += 1


def _render_code_block(pdf, code_lines, effective_width):
    """渲染代码块：Courier 字体 + 灰色背景。"""
    if not code_lines:
        return
    pdf.set_font("Courier", "", 8)
    pdf.set_fill_color(245, 245, 245)
    for cl in code_lines:
        safe_line = cl[:130] if len(cl) > 130 else cl
        try:
            pdf.cell(effective_width, 5, safe_line, fill=True)
            pdf.ln(5)
        except Exception:
            pdf.ln(4)
    pdf.ln(2)


def _render_table(pdf, font_name, table_lines, effective_width):
    """渲染简单 Markdown 表格。

    使用单元格边框和居中对齐，表头加粗显示。
    """
    try:
        rows = []
        for tl in table_lines:
            cells = [c.strip() for c in tl.strip().split("|")][1:-1]  # 去除首尾空
            if cells:
                rows.append(cells)

        if len(rows) < 2:
            return

        cols = max(len(r) for r in rows)
        col_w = effective_width / max(cols, 1)

        pdf.set_font(font_name, "", 8)
        pdf.set_fill_color(230, 235, 245)

        for ri, row in enumerate(rows):
            # 填充缺失列
            while len(row) < cols:
                row.append("")

            if ri == 0:
                pdf.set_font(font_name, "B", 9)
            else:
                pdf.set_font(font_name, "", 8)

            for ci, cell in enumerate(row):
                cell_text = cell[:40] if len(cell) > 40 else cell
                if ri == 0:
                    pdf.cell(col_w, 7, cell_text, border=1, fill=True, align="C")
                else:
                    pdf.cell(col_w, 7, cell_text, border=1, align="C")
            pdf.ln(7)

        pdf.ln(3)
    except Exception:
        pdf.ln(3)
        pdf.set_font(font_name, "", 9)
        pdf.multi_cell(effective_width, 6, "[表格渲染失败]")


def _render_list_item(pdf, font_name, line, effective_width):
    """渲染列表项。保留公式和格式标记。"""
    try:
        pdf.set_font(font_name, "", 10)
        text = line.strip()
        # 提取列表标记
        marker_match = re.match(r"^(\s*[-*+]\s+|\s*\d+[.)]\s+)", text)
        marker = marker_match.group(0) if marker_match else "• "
        text = text[len(marker):] if marker_match else text
        pdf.cell(5, 6, "•")
        pdf.multi_cell(effective_width - 5, 6, text)
    except Exception:
        pdf.ln(4)


def _render_blockquote(pdf, font_name, line, effective_width):
    """渲染引用块：左侧竖线 + 缩进。"""
    try:
        pdf.set_font(font_name, "I", 10)
        text = line.strip()[2:]  # 去掉 "> "
        pdf.set_x(pdf.l_margin + 6)
        pdf.set_draw_color(100, 100, 100)
        pdf.line(pdf.l_margin + 2, pdf.get_y(), pdf.l_margin + 2, pdf.get_y() + 6)
        pdf.multi_cell(effective_width - 8, 6, text)
        pdf.set_draw_color(0, 0, 0)
    except Exception:
        pdf.ln(4)


def _render_paragraph(pdf, font_name, line, effective_width):
    """渲染普通段落。保留 LaTeX 公式（$...$ / $$...$$）、粗体斜体、链接。"""
    try:
        pdf.set_font(font_name, "", 10)
        text = line.strip()
        if not text:
            return

        # 保留公式和链接格式，仅清理不可渲染字符
        text = text.replace("\t", "    ")

        if len(text) > 200:
            # 长段落分多次渲染
            pdf.multi_cell(effective_width, 6, text)
        else:
            pdf.multi_cell(effective_width, 6, text)
    except Exception:
        pdf.ln(4)


def _pdf_section_heading(pdf, font_name: str, heading: str, level: int):
    """渲染 PDF 章节标题。"""
    sizes = {1: ("B", 16), 2: ("B", 14), 3: ("B", 12), 4: ("", 11)}
    style, size = sizes.get(level, ("", 11))
    pdf.ln(4)
    pdf.set_font(font_name, style, size)
    try:
        pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, size * 0.6, heading)
    except Exception:
        pdf.cell(0, size * 0.6, heading[:50])
    pdf.ln(2)


# ============================================================
# Word 生成
# ============================================================

def _generate_docx(content: str, title: str, author: str,
                   folder: Path, safe_title: str, timestamp: str) -> Path:
    """使用 python-docx 生成 Word 文档。保留公式和格式。"""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.core_properties.title = title
    if author:
        doc.core_properties.author = author

    # 标题页
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(22)
    run.font.bold = True
    doc.add_paragraph()

    if author:
        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = author_para.add_run(f"作者: {author}")
        run.font.size = Pt(12)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"生成时间: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC")
    run.font.size = Pt(10)
    doc.add_paragraph()
    doc.add_paragraph()

    _render_markdown_to_docx(doc, content)

    filepath = folder / f"{safe_title}_{timestamp}.docx"
    doc.save(str(filepath))
    return filepath


def _render_markdown_to_docx(doc, content: str):
    """将 Markdown 渲染为 Word 段落。保留 LaTeX 公式。"""
    from docx.shared import Pt

    lines = content.splitlines()
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code_block:
                for cl in code_lines:
                    code_para = doc.add_paragraph()
                    run = code_para.add_run(cl)
                    run.font.name = "Courier New"
                    run.font.size = Pt(8)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 表格
        if "|" in line and line.strip().startswith("|"):
            table_lines = [line]
            j = i + 1
            while j < len(lines) and "|" in lines[j] and lines[j].strip().startswith("|"):
                table_lines.append(lines[j])
                j += 1
            filtered = [tl for tl in table_lines if not re.match(r"^\|[\s\-:|]+\|$", tl.strip())]
            if len(filtered) >= 2:
                _render_docx_table(doc, filtered)
            else:
                doc.add_paragraph(line.strip())
            i = j
            continue

        # 标题
        h_match = re.match(r"^(#{1,4})\s+(.+)", line)
        if h_match:
            level = len(h_match.group(1))
            heading_text = h_match.group(2).strip()
            heading_style = f"Heading {min(level, 4)}"
            try:
                para = doc.add_paragraph()
                para.style = doc.styles[heading_style]
                para.add_run(heading_text)
            except KeyError:
                para = doc.add_paragraph()
                run = para.add_run(heading_text)
                run.font.bold = True
                run.font.size = Pt(18 - level * 2)
            i += 1
            continue

        # 列表
        if re.match(r"^\s*[-*+]\s+", line):
            text = line.strip()[2:] if line.strip().startswith("- ") else line.strip()[1:]
            try:
                doc.add_paragraph(text, style="List Bullet")
            except KeyError:
                para = doc.add_paragraph()
                para.add_run(f"  • {text}")
            i += 1
            continue

        # 有序列表
        if re.match(r"^\s*\d+[.)]\s+", line):
            try:
                doc.add_paragraph(line.strip(), style="List Number")
            except KeyError:
                doc.add_paragraph(line.strip())
            i += 1
            continue

        # 水平线
        if re.match(r"^[-*_]{3,}\s*$", line):
            para = doc.add_paragraph()
            para.add_run("─" * 60)
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 普通段落
        if line.strip():
            doc.add_paragraph(line.strip())
        i += 1


def _render_docx_table(doc, table_lines):
    """在 Word 文档中渲染表格。"""
    from docx.shared import Pt

    rows = []
    for tl in table_lines:
        cells = [c.strip() for c in tl.strip().split("|")][1:-1]
        if cells:
            rows.append(cells)

    if len(rows) < 2:
        return

    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols, style="Table Grid")

    for ri, row_data in enumerate(rows):
        while len(row_data) < cols:
            row_data.append("")
        for ci, cell_text in enumerate(row_data):
            cell = table.cell(ri, ci)
            cell.text = cell_text
            if ri == 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
    doc.add_paragraph()


# ============================================================
# Markdown 生成
# ============================================================

def _generate_markdown(content: str, folder: Path,
                       safe_title: str, timestamp: str) -> Path:
    """直接保存 Markdown 文本文件。"""
    filepath = folder / f"{safe_title}_{timestamp}.md"
    filepath.write_text(content, encoding="utf-8")
    return filepath

"""
文档解析模块 —— 提取图片、PDF、Word、Markdown 文档中的数学题目。

支持的文件类型：
- 图片 (png/jpg/jpeg/webp): 转换为 Base64 供视觉模型识别
- PDF: 使用 pdfplumber 提取文本内容
- Word (docx): 使用 python-docx 提取段落文本
- Markdown (md): 直接读取文本内容

使用方式：
    from utils.document_parser import parse_document
    text, image_b64 = parse_document(file_path, file_type)
"""

import base64
import io
from pathlib import Path
from typing import Optional, Tuple


def parse_document(file_path: str, file_type: str) -> Tuple[str, Optional[str]]:
    """解析文档文件，返回 (提取文本, 图片Base64或None)。

    Args:
        file_path: 文件绝对路径
        file_type: 文件类型标识 ("image" / "pdf" / "docx" / "markdown")

    Returns:
        (提取的文本内容, 图片Base64字符串或None)
        对于图片文件，文本为提示词，image_b64包含Base64编码的图片数据
        对于其他文件，文本为提取的内容，image_b64为None
    """
    ext = Path(file_path).suffix.lower()

    if file_type == "image" or ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp"}:
        return _parse_image(file_path)
    elif file_type == "pdf" or ext == ".pdf":
        return _parse_pdf(file_path), None
    elif file_type == "docx" or ext in {".docx", ".doc"}:
        return _parse_docx(file_path), None
    elif file_type == "markdown" or ext == ".md":
        return _parse_markdown(file_path), None
    else:
        return f"[无法解析的文件类型: {ext}]", None


def _parse_image(file_path: str) -> Tuple[str, Optional[str]]:
    """将图片文件编码为 Base64 Data URI，供视觉模型直接识别。

    识别提示词要求模型提取图片中的数学题目和公式。

    Args:
        file_path: 图片文件路径

    Returns:
        (识别提示词, Base64 Data URI)
    """
    with open(file_path, "rb") as f:
        img_bytes = f.read()

    # 检测 MIME 类型
    ext = Path(file_path).suffix.lower()
    mime_map = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".bmp": "image/bmp",
        ".gif": "image/gif",
    }
    mime_type = mime_map.get(ext, "image/png")

    b64_data = base64.b64encode(img_bytes).decode("ascii")
    data_uri = f"data:{mime_type};base64,{b64_data}"

    prompt = (
        "请仔细查看这张图片，识别并提取其中所有的数学题目、公式和图表信息。"
        "如果是手写内容，请尽可能准确地进行OCR识别。"
        "将识别结果以清晰的文本格式输出，保留数学公式的LaTeX表示。"
    )

    return prompt, data_uri


def _parse_pdf(file_path: str) -> str:
    """从 PDF 文件中提取文本内容。

    使用 pdfplumber 提取，回退到 PyPDF2。
    对提取的文本进行基本清洗，去除多余空行。

    Args:
        file_path: PDF 文件路径

    Returns:
        提取的文本内容，失败时返回错误提示。
    """
    text_parts = []

    # 优先使用 pdfplumber（提取质量更好）
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"--- 第{i+1}页 ---\n{page_text.strip()}")
        if text_parts:
            result = "\n\n".join(text_parts)
            if len(result) > 8000:
                result = result[:8000] + "\n...[内容过长已截断]"
            return result
    except Exception as e:
        text_parts.append(f"[pdfplumber解析失败: {e}]")

    # 回退：PyPDF2
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- 第{i+1}页 ---\n{page_text.strip()}")
        if text_parts:
            result = "\n\n".join(text_parts)
            if len(result) > 8000:
                result = result[:8000] + "\n...[内容过长已截断]"
            return result
    except Exception as e:
        text_parts.append(f"[PyPDF2解析失败: {e}]")

    if not text_parts:
        return "[PDF解析失败] 无法从此PDF文件中提取文本。请检查文件是否为扫描图片版PDF，建议转换为文本格式后再试。"
    return "\n\n".join(text_parts)


def _parse_docx(file_path: str) -> str:
    """从 Word (.docx) 文件中提取文本内容。

    提取所有段落文本和表格内容。

    Args:
        file_path: .docx 文件路径

    Returns:
        提取的文本内容。
    """
    try:
        from docx import Document
        doc = Document(file_path)

        text_parts = []

        # 提取段落文本
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        result = "\n\n".join(text_parts)
        if not result:
            return "[Word文档为空] 未找到可提取的文本内容。"

        if len(result) > 8000:
            result = result[:8000] + "\n...[内容过长已截断]"
        return result

    except Exception as e:
        return f"[Word解析失败: {e}]"


def _parse_markdown(file_path: str) -> str:
    """读取 Markdown 文件的文本内容。

    Args:
        file_path: .md 文件路径

    Returns:
        文件文本内容（最多 6000 字符）。
    """
    try:
        text = Path(file_path).read_text(encoding="utf-8", errors="ignore")
        if len(text) > 6000:
            text = text[:6000] + "\n...[内容过长已截断]"
        return text
    except Exception as e:
        return f"[Markdown读取失败: {e}]"


def get_mime_type(file_path: str) -> str:
    """根据文件扩展名获取 MIME 类型。"""
    ext = Path(file_path).suffix.lower()
    return {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".webp": "image/webp",
        ".gif": "image/gif",
        ".pdf": "application/pdf",
        ".doc": "application/msword",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".md": "text/markdown",
    }.get(ext, "application/octet-stream")
